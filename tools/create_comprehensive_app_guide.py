from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
PORTFOLIO_JSON = ROOT / "website" / "public" / "data" / "portfolio.json"
SYNC_CONFIG = ROOT / "tools" / "github_sync_config.json"


NAVY = RGBColor(5, 31, 57)
TEAL = RGBColor(0, 132, 132)
GOLD = RGBColor(190, 141, 45)
GREY = RGBColor(236, 240, 243)
TEXT = RGBColor(31, 43, 55)


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8-sig"))


def money(value) -> str:
    try:
        return f"EGP {float(value):,.0f}"
    except Exception:
        return "N/A"


def pct(value) -> str:
    try:
        return f"{float(value) * 100:.1f}%"
    except Exception:
        return "N/A"


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    lines += ["| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |" for row in rows]
    return "\n".join(lines)


def build_markdown(portfolio: dict, sync_config: dict) -> str:
    projects = portfolio.get("projects", [])
    sectors = portfolio.get("sectors", [])
    totals = portfolio.get("totals", {})
    generated = datetime.now().strftime("%Y-%m-%d %H:%M")

    project_rows = [
        [
            p.get("project_display_name", "N/A"),
            p.get("sector", "N/A"),
            p.get("status", "N/A"),
            money(p.get("contract_value")),
            pct(p.get("actual_progress")),
            str(p.get("activity_count", "N/A")),
            "Yes" if p.get("decision_required") else "No",
        ]
        for p in projects
    ]

    sector_rows = [
        [
            s.get("sector", "N/A"),
            str(s.get("project_count", "N/A")),
            money(s.get("contract_value")),
            pct(s.get("average_progress")),
            str(round(s.get("average_spi") or 0, 2)),
            str(round(s.get("average_cpi") or 0, 2)),
            str(s.get("decisions_required", "N/A")),
        ]
        for s in sectors
    ]

    tab_rows = [
        ["Decision Making Dashboard", "Top-management portfolio view", "Aggregates all project JSON into portfolio KPIs, sector view, project console, and report viewer."],
        ["Overview", "Project executive summary", "Shows project status, value, progress, remaining value, source metadata, and overview table."],
        ["EBS / WBS", "Work breakdown view", "Shows WBS source table and master dashboard embed for the selected project."],
        ["Activities", "Activity register", "Shows activity, progress, EVM, and delay event counts with source-table preview."],
        ["Milestones", "Milestone tracking", "Shows milestone records, SPI schedule health, delay days, and milestone source table."],
        ["S-Curve", "Progress trend", "Shows S-curve source status and linked executive dashboard HTML."],
        ["EVM", "Earned value management", "Calculates BAC, PV, EV, AC, SPI, CPI from selected project data."],
        ["Contracts", "Commercial position", "Shows contract value, paid, spent, remaining value, contracts and payments previews."],
        ["Risk Matrix", "Risk and decision signals", "Shows risk score, risk rows, decision trigger, delay exposure, and risk table."],
        ["Letters Intelligence", "Correspondence intelligence", "Detects inbox files and workbook, shows letter inventory and generated dashboard output."],
        ["Delay Analysis", "Time Impact Analysis", "Detects TIA templates, missing files, column previews, MEP schedule, BL schedule, and generated delay charts."],
        ["Contract & Claims", "Claims knowledge base", "Detects contract files, clause library, SQLite tables, evidence, and claim exposure."],
        ["Conference", "Meeting support", "Reads meeting_url from project.json and provides same-page conference panel and join button."],
        ["Output Studio", "Generated reports", "Shows four automatic HTML outputs and watcher/sync status for the selected project."],
    ]

    folder_rows = [
        ["projects/{Sector}/{Project}", "Main project workspace", "Add future projects here under sector folders. The generator detects them automatically."],
        ["01-data/import_templates", "Core project data", "CSV files for projects, activities, progress, EVM, payments, contracts, claims, risks, milestones, WBS, S-curve."],
        ["02-delay_analysis", "TIA workspace", "Methodology documents and steel_delay_tia_templates CSVs used by Delay Analysis."],
        ["03-schedule", "Schedule support", "BL Schedule, MEP Schedule, MEP Activities, and civil logic files."],
        ["04-source_excel", "Raw Excel source", "Original Excel files before conversion into controlled CSV inputs."],
        ["05-contracts", "Contract intelligence", "Contract source PDFs, clause libraries, contract_claims.db, and clauses."],
        ["06-evidence", "Evidence register", "Evidence templates, photo logs, document references, and claim evidence."],
        ["07-letters_intelligence", "Correspondence intelligence", "letters_intelligence.xlsx plus inbox folders for new letters."],
        ["08-branding", "Project branding", "Logo placeholder, identity, palette, and report branding templates."],
        ["09-notes", "Project notes", "Meeting, engineering, and claims notes."],
        ["10-deliverables", "Formal deliverables", "Generated project deliverables outside the automatic website HTML output set."],
        ["11-outputs", "Automatic HTML outputs", "Per-project generated HTML outputs: executive, master, SVG charts, linked dashboard."],
        ["12-logs", "Project logs", "Project-specific operational logs."],
        ["website", "Next.js public site", "Vercel website source, React UI, CSS, and public data."],
        ["tools", "Automation tools", "Data generator, GitHub no-Git sync, validation, lineage, and package builders."],
        ["src/construction_system", "Python backend logic", "Streamlit support modules, project context, loaders, TIA logic, reporting, OpenAI gateway."],
        ["dashboard.py", "Original Streamlit app", "The full original application with Python-based calculations and advanced workflows."],
    ]

    output_rows = [
        ["01_executive_dashboard.html", "Executive dashboard", "Project-level management dashboard HTML."],
        ["02_master_dashboard.html", "Master dashboard", "Detailed project dashboard by sections."],
        ["03_elite_svg_charts.html", "Elite SVG charts", "Tabbed SVG chart gallery generated as HTML."],
        ["04_linked_executive_dashboard.html", "Linked executive dashboard", "Linked A3-style executive dashboard HTML."],
    ]

    commands = [
        ["Generate website data", 'cd "D:\\Project Intelligence Hub NextJS" && python tools\\generate_nextjs_website_data.py'],
        ["Run website locally", 'cd "D:\\Project Intelligence Hub NextJS\\website" && npm run dev'],
        ["Build website", 'cd "D:\\Project Intelligence Hub NextJS\\website" && npm run build'],
        ["Deploy to Vercel", 'cd "D:\\Project Intelligence Hub NextJS\\website" && npx vercel@latest deploy --prod --yes --project samcoegyptdashboard'],
        ["Sync to GitHub once", 'cd "D:\\Project Intelligence Hub NextJS" && cmd /c RUN_FULL_PROJECT_NO_GIT_SYNC.bat Once 30'],
        ["Start sync watcher", 'cd "D:\\Project Intelligence Hub NextJS" && cmd /c RUN_FULL_PROJECT_NO_GIT_SYNC.bat Watch 30'],
    ]

    return f"""# Project Intelligence Hub - Comprehensive App Guide

Generated: {generated}

Live website: https://samcoegyptdashboard.vercel.app

Figma reference design: https://www.figma.com/design/z1T4ERNeuBlrpC3slf93KT

## 1. Executive Summary

Project Intelligence Hub is a project-controls, delay-analysis, claims-intelligence, and executive-reporting system. It currently has two working surfaces:

- Original Streamlit app: the full Python application, advanced calculations, uploads, exports, and detailed project-controls workflows.
- Next.js/Vercel website: public, fast, mobile-friendly executive website generated from the same project folders and HTML outputs.

The main operating rule is project isolation: each selected project reads from its own project folder, shows its own data, and writes or publishes its own outputs.

## 2. Current Portfolio Snapshot

- Projects detected: {portfolio.get("project_count", 0)}
- Sectors detected: {portfolio.get("sector_count", 0)}
- Total contract value: {money(totals.get("contract_value"))}
- Average progress: {pct(totals.get("average_progress"))}
- Average SPI: {round(totals.get("average_spi") or 0, 2)}
- Average CPI: {round(totals.get("average_cpi") or 0, 2)}
- Decisions required: {totals.get("decisions_required", 0)}

{md_table(["Project", "Sector", "Status", "Contract Value", "Progress", "Activities", "Decision Required"], project_rows)}

## 3. Sector Structure

Projects are grouped by sector folders under `projects`. Adding a new sector folder automatically creates a new sector in the Decision Making Dashboard after the generator runs.

{md_table(["Sector", "Projects", "Contract Value", "Progress", "SPI", "CPI", "Decisions"], sector_rows)}

## 4. High-Level Architecture

```mermaid
flowchart LR
    A["Project folders"] --> B["Python data generator"]
    B --> C["website/public/data JSON"]
    B --> D["website/public/generated HTML outputs"]
    C --> E["Next.js website"]
    D --> E
    E --> F["Vercel public URL"]
    A --> G["Streamlit original app"]
    G --> H["Advanced reports, TIA, claims, exports"]
    I["RUN_FULL_PROJECT_NO_GIT_SYNC.bat"] --> J["GitHub repository"]
    J --> F
```

## 5. Main User Flow

1. Open the website.
2. Decision Making Dashboard appears first.
3. Choose a project from the Active Project dropdown.
4. The same page updates the Project Workspace tabs.
5. Use the old-app-style tabs for Overview, EBS/WBS, Activities, Milestones, S-Curve, EVM, Contracts, Risk Matrix, Letters Intelligence, Delay Analysis, Contract & Claims, Conference, and Output Studio.
6. Update source files in the selected project folder.
7. Run the website data generator and sync.
8. Vercel updates the public dashboard from GitHub.

## 6. App Tabs and Features

{md_table(["Tab / Page", "Purpose", "Main Data or Behavior"], tab_rows)}

## 7. Folder Guide

{md_table(["Folder / File Area", "Purpose", "How It Is Used"], folder_rows)}

## 8. Project Folder Standard

Every project should live at:

`projects/{{Sector Name}}/{{Project Folder Name}}`

Minimum files:

- `project.json`
- `project_manifest.json`
- `01-data/import_templates/*.csv`
- `02-delay_analysis/steel_delay_tia_templates/*.csv`
- `03-schedule/*.csv`
- `05-contracts/source/*`
- `07-letters_intelligence/inbox/*`

Recommended `project.json` fields:

```json
{{
  "project_id": "unique-project-id",
  "project_name": "Visible Project Name",
  "client_name": "Client",
  "contractor": "SAMCO",
  "currency": "EGP",
  "status": "Active",
  "meeting_url": "https://your-meeting-link"
}}
```

## 9. Data Pipeline

The website data pipeline is:

Project folders -> `tools/generate_nextjs_website_data.py` -> `website/public/data/portfolio.json` and `website/public/data/projects/*.json` -> Next.js UI.

The generator also copies automatic HTML outputs from root `11-outputs/{{Project Folder Name}}` into `website/public/generated/{{Project Folder Name Slug}}`.

## 10. Automatic HTML Outputs

Each project should have exactly these automatic website HTML outputs:

{md_table(["File", "Name", "Purpose"], output_rows)}

## 11. Letters Intelligence

Letters Intelligence is project-scoped. The website detects:

- `07-letters_intelligence/letters_intelligence.xlsx`
- `07-letters_intelligence/inbox/**`
- letter file count, extension, size, modified date
- linked claims and delay-event availability

To add new letters, place them in the selected project inbox folder, then regenerate and sync.

## 12. Delay Analysis - Time Impact Analysis

Delay Analysis reads project-specific TIA templates from:

`02-delay_analysis/steel_delay_tia_templates`

The website detects the file inventory, row count, column count, missing required files, MEP activities, MEP schedule, MEP civil logic, and BL schedule. The original Streamlit app remains the full calculation engine for deep TIA exports and report generation.

## 13. Contract & Claims Intelligence

Contract & Claims uses:

- `05-contracts/source`
- `05-contracts/source/Overall_Contract_clause_library.xlsx`
- `05-contracts/contract_claims.db`
- `06-evidence`

The website shows file inventory and knowledge-base table counts. The original Streamlit app remains the advanced claims engine for extraction, evidence mapping, rebuttal, claim drafting, and exports.

## 14. Conference Call

The Conference tab reads `meeting_url` from the selected project's `project.json`.

If the field is missing, the tab shows setup instructions. If the field exists, it shows a Join Conference button.

Teams and Zoom usually block iframe embedding, so the meeting may open in a new tab while the dashboard remains available.

## 15. Output Studio

Output Studio shows the four automatic HTML reports and sync/watch status. Manual Streamlit output features still exist in the original app. The Next.js website focuses on fast HTML report viewing and mobile-safe access.

## 16. Sync and GitHub

The no-Git sync is controlled by:

- `RUN_FULL_PROJECT_NO_GIT_SYNC.bat`
- `tools/github_no_git_sync.ps1`
- `tools/github_sync_config.json`

Repository target:

- Owner: {sync_config.get("owner", "N/A")}
- Repository: {sync_config.get("repository", "N/A")}
- Branch: {sync_config.get("branch", "N/A")}
- Default interval seconds: {sync_config.get("interval_seconds", "N/A")}
- Deletion sync: {sync_config.get("sync_deletions", "N/A")}

Credentials must come from `GITHUB_TOKEN` or `GH_TOKEN` in the local Windows process.

## 17. Vercel Deployment

Recommended Vercel settings:

- Root directory: `website`
- Framework: Next.js
- Build command: `npm run build`
- Install command: `npm install`
- Output directory: `.next`

Production URL:

`https://samcoegyptdashboard.vercel.app`

## 18. Common Commands

{md_table(["Action", "Command"], commands)}

## 19. Adding a New Project

1. Create a sector folder under `projects` if needed.
2. Copy `_PROJECT_TEMPLATE` into the sector folder.
3. Rename the copied folder to the project name.
4. Edit `project.json` and `project_manifest.json`.
5. Add project CSVs under `01-data/import_templates`.
6. Add TIA files under `02-delay_analysis/steel_delay_tia_templates`.
7. Add schedule support under `03-schedule`.
8. Add contracts under `05-contracts/source`.
9. Add letters under `07-letters_intelligence/inbox`.
10. Run the generator.
11. Build and sync.

## 20. Update Workflow

When project files change:

1. Update the files in the correct project folder.
2. Run `python tools\\generate_nextjs_website_data.py`.
3. Run `npm run build` from `website`.
4. Run `RUN_FULL_PROJECT_NO_GIT_SYNC.bat Once 30`.
5. Vercel rebuilds from GitHub.

## 21. Troubleshooting

| Issue | Likely Cause | Fix |
| --- | --- | --- |
| Project not visible | Folder not under sector folder or generator not run | Place folder under `projects/{{Sector}}` and rerun generator |
| Website shows old data | JSON not regenerated or GitHub/Vercel not updated | Run generator, sync, redeploy |
| Conference button missing | `meeting_url` missing in `project.json` | Add meeting URL and regenerate |
| Letters not detected | Files not in project inbox | Add files under `07-letters_intelligence/inbox` |
| Delay tab missing files | Required TIA CSVs absent or named differently | Add files to `02-delay_analysis/steel_delay_tia_templates` |
| Claims data looks empty | Contract source or DB missing | Add contract source files and rebuild contract library in Streamlit |
| Sync fails 401 | Token not available in process scope | Set `GITHUB_TOKEN` or `GH_TOKEN` in the active Windows environment |
| Vercel build warning about audit | Dev-package audit noise | Check production audit with `npm audit --omit=dev` |

## 22. Maintenance Rules

- Do not mix project data between folders.
- Do not hardcode project names in the website UI.
- Use `project.json` for project identity and meeting links.
- Use CSV files as the source of visible project data.
- Regenerate website JSON after source changes.
- Keep automatic outputs as HTML in per-project output folders.
- Keep credentials outside the repository.
- Use the Streamlit app for deep Python workflows and the Next.js site for public/mobile viewing.

## 23. Roadmap

Recommended next improvements:

1. Add a backend API if live online uploads and AI extraction are required on Vercel.
2. Add project-level authentication if dashboard access must be private.
3. Add scheduled GitHub/Vercel automation for automatic rebuilds.
4. Add Playwright tests for tab clicks and mobile rendering.
5. Add a small admin page to edit `project.json` safely.
6. Add export status badges for each generated report.
"""


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = NAVY if level == 1 else TEAL
        run.font.name = "Arial"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "ECF0F3")
        shading.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = NAVY
                run.font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = TEXT


def write_docx(markdown: str, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = TEXT

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Project Intelligence Hub")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Comprehensive App Guide")
    r.font.name = "Arial"
    r.font.size = Pt(14)
    r.font.color.rgb = GOLD
    doc.add_paragraph()

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:], 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Inches(0.2)
            p.add_run("• ").font.color.rgb = TEAL
            p.add_run(line[2:])
        elif line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(60, 72, 82)
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            headers = [cell.strip() for cell in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([cell.strip() for cell in lines[i].strip("|").split("|")])
                i += 1
            add_table(doc, headers, rows)
            continue
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(4)
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Project Intelligence Hub | Comprehensive App Guide | Confidential")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)
    doc.save(output)


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    portfolio = load_json(PORTFOLIO_JSON)
    sync_config = load_json(SYNC_CONFIG)
    markdown = build_markdown(portfolio, sync_config)
    md_path = REPORTS / "Project_Intelligence_Hub_Comprehensive_Guide.md"
    docx_path = REPORTS / "Project_Intelligence_Hub_Comprehensive_Guide.docx"
    md_path.write_text(markdown, encoding="utf-8")
    write_docx(markdown, docx_path)
    print(md_path)
    print(docx_path)


if __name__ == "__main__":
    main()
